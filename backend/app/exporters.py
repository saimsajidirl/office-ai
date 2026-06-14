import csv
import json
import re
from io import BytesIO, StringIO
from typing import Callable
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from .schemas import ConversionPlan, OutputFormat, TableData


MIME_TYPES = {
    OutputFormat.CSV: "text/csv; charset=utf-8",
    OutputFormat.XLSX: (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    ),
    OutputFormat.DOCX: (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ),
    OutputFormat.PDF: "application/pdf",
    OutputFormat.JSON: "application/json",
    OutputFormat.TXT: "text/plain; charset=utf-8",
}


def safe_filename(title: str) -> str:
    clean = re.sub(r"[^A-Za-z0-9_-]+", "-", title.strip()).strip("-").lower()
    return clean[:60] or "converted-data"


def safe_spreadsheet_value(value: str) -> str:
    if value.startswith(("=", "+", "-", "@")):
        return f"'{value}"
    return value


def pdf_text(value: str) -> str:
    return escape(str(value)).replace("\n", "<br/>")


def normalized_table(plan: ConversionPlan) -> TableData:
    columns = [
        str(column).strip() or f"Column {index + 1}"
        for index, column in enumerate(plan.table.columns)
    ]
    rows = [[str(value).strip() for value in row] for row in plan.table.rows]

    if not columns:
        columns = ["Content"]
        rows = [[plan.summary]]

    width = len(columns)
    normalized_rows = [
        (row + [""] * width)[:width]
        for row in rows
        if any(value for value in row)
    ]
    return TableData(
        title=plan.table.title or plan.title,
        columns=columns,
        rows=normalized_rows,
    )


def render_csv(plan: ConversionPlan) -> bytes:
    table = normalized_table(plan)
    stream = StringIO(newline="")
    writer = csv.writer(stream)
    writer.writerow(safe_spreadsheet_value(value) for value in table.columns)
    writer.writerows(
        [safe_spreadsheet_value(value) for value in row] for row in table.rows
    )
    return stream.getvalue().encode("utf-8-sig")


def render_xlsx(plan: ConversionPlan) -> bytes:
    table = normalized_table(plan)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Converted Data"
    sheet.freeze_panes = "A2"
    sheet.append(table.columns)

    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="4338CA")
        cell.alignment = Alignment(horizontal="center")

    for row in table.rows:
        sheet.append([safe_spreadsheet_value(value) for value in row])

    for column_index, column in enumerate(table.columns, start=1):
        values = [column] + [
            row[column_index - 1] for row in table.rows if len(row) >= column_index
        ]
        width = min(max(len(str(value)) for value in values) + 2, 48)
        sheet.column_dimensions[get_column_letter(column_index)].width = max(width, 12)

    stream = BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def render_docx(plan: ConversionPlan) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    title = document.add_heading(plan.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(plan.summary)

    for content_section in plan.sections:
        if content_section.heading:
            document.add_heading(content_section.heading, level=1)
        for paragraph in content_section.paragraphs:
            document.add_paragraph(paragraph)
        for bullet in content_section.bullets:
            document.add_paragraph(bullet, style="List Bullet")

    table_data = normalized_table(plan)
    if table_data.rows:
        document.add_heading(table_data.title, level=1)
        table = document.add_table(rows=1, cols=len(table_data.columns))
        table.style = "Table Grid"
        for index, column in enumerate(table_data.columns):
            cell = table.rows[0].cells[index]
            cell.text = column
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        for row in table_data.rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value

    if plan.notes:
        document.add_heading("Notes", level=1)
        for note in plan.notes:
            document.add_paragraph(note, style="List Bullet")

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def render_pdf(plan: ConversionPlan) -> bytes:
    stream = BytesIO()
    document = SimpleDocTemplate(
        stream,
        pagesize=landscape(A4),
        leftMargin=0.45 * inch,
        rightMargin=0.45 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
        title=plan.title,
    )
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CenteredTitle",
            parent=styles["Title"],
            alignment=TA_CENTER,
            textColor=colors.HexColor("#312E81"),
        )
    )
    story = [
        Paragraph(pdf_text(plan.title), styles["CenteredTitle"]),
        Paragraph(pdf_text(plan.summary), styles["BodyText"]),
        Spacer(1, 12),
    ]

    for content_section in plan.sections:
        if content_section.heading:
            story.append(
                Paragraph(pdf_text(content_section.heading), styles["Heading2"])
            )
        story.extend(
            Paragraph(pdf_text(paragraph), styles["BodyText"])
            for paragraph in content_section.paragraphs
        )
        if content_section.bullets:
            story.append(
                ListFlowable(
                    [
                        ListItem(Paragraph(pdf_text(item), styles["BodyText"]))
                        for item in content_section.bullets
                    ],
                    bulletType="bullet",
                )
            )
        story.append(Spacer(1, 8))

    table_data = normalized_table(plan)
    if table_data.rows:
        story.append(Paragraph(pdf_text(table_data.title), styles["Heading2"]))
        pdf_rows = [
            [
                Paragraph(pdf_text(value), styles["BodyText"])
                for value in table_data.columns
            ]
        ]
        pdf_rows.extend(
            [Paragraph(pdf_text(value), styles["BodyText"]) for value in row]
            for row in table_data.rows
        )
        page_width = landscape(A4)[0] - 0.9 * inch
        column_width = page_width / max(len(table_data.columns), 1)
        table = Table(
            pdf_rows,
            colWidths=[column_width] * len(table_data.columns),
            repeatRows=1,
        )
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4338CA")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    (
                        "ROWBACKGROUNDS",
                        (0, 1),
                        (-1, -1),
                        [colors.white, colors.HexColor("#F8FAFC")],
                    ),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.extend([table, Spacer(1, 10)])

    if plan.notes:
        story.append(Paragraph("Notes", styles["Heading2"]))
        story.append(
            ListFlowable(
                [
                    ListItem(Paragraph(pdf_text(note), styles["BodyText"]))
                    for note in plan.notes
                ],
                bulletType="bullet",
            )
        )

    document.build(story)
    return stream.getvalue()


def render_json(plan: ConversionPlan) -> bytes:
    table = normalized_table(plan)
    records = [dict(zip(table.columns, row, strict=False)) for row in table.rows]
    payload = {
        "title": plan.title,
        "summary": plan.summary,
        "records": records,
        "notes": plan.notes,
    }
    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")


def render_txt(plan: ConversionPlan) -> bytes:
    lines = [plan.title, "=" * len(plan.title), "", plan.summary, ""]
    for section in plan.sections:
        if section.heading:
            lines.extend([section.heading, "-" * len(section.heading)])
        lines.extend(section.paragraphs)
        lines.extend(f"- {bullet}" for bullet in section.bullets)
        lines.append("")

    table = normalized_table(plan)
    if table.rows:
        lines.extend([table.title, "-" * len(table.title)])
        lines.append("\t".join(table.columns))
        lines.extend("\t".join(row) for row in table.rows)
        lines.append("")

    if plan.notes:
        lines.extend(["Notes", "-----"])
        lines.extend(f"- {note}" for note in plan.notes)

    return "\n".join(lines).strip().encode("utf-8")


EXPORTERS: dict[OutputFormat, Callable[[ConversionPlan], bytes]] = {
    OutputFormat.CSV: render_csv,
    OutputFormat.XLSX: render_xlsx,
    OutputFormat.DOCX: render_docx,
    OutputFormat.PDF: render_pdf,
    OutputFormat.JSON: render_json,
    OutputFormat.TXT: render_txt,
}


def render_file(plan: ConversionPlan, output_format: OutputFormat) -> tuple[bytes, str, str]:
    content = EXPORTERS[output_format](plan)
    filename = f"{safe_filename(plan.title)}.{output_format.value}"
    return content, filename, MIME_TYPES[output_format]


def create_preview(plan: ConversionPlan, output_format: OutputFormat) -> str:
    if output_format in {
        OutputFormat.CSV,
        OutputFormat.XLSX,
        OutputFormat.JSON,
    }:
        table = normalized_table(plan)
        lines = [" | ".join(table.columns)]
        lines.extend(" | ".join(row) for row in table.rows[:8])
        if len(table.rows) > 8:
            lines.append(f"... and {len(table.rows) - 8} more rows")
        return "\n".join(lines)

    lines = [plan.title, plan.summary]
    for section in plan.sections[:3]:
        lines.append(section.heading)
        lines.extend(section.paragraphs[:2])
        lines.extend(f"- {item}" for item in section.bullets[:3])
    return "\n\n".join(value for value in lines if value)[:4_000]
